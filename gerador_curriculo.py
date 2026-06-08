"""
Módulo de geração de currículo em PDF e DOCX
Baseado no modelo do currículo do Adriano Pontes Dias
"""

import io
from PIL import Image

# ── ReportLab (PDF) ──────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
)
from reportlab.platypus import Image as RLImage
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib.utils import ImageReader

# ── python-docx (Word) ───────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

# ── Cores do tema ─────────────────────────────────────────────────────────────
COR_PRIMARIA   = colors.HexColor("#1a3c5e")   # azul marinho
COR_SECUNDARIA = colors.HexColor("#2e75b6")   # azul médio
COR_LINHA      = colors.HexColor("#d0dce8")
COR_FUNDO      = colors.HexColor("#f0f5fa")
COR_TEXTO      = colors.HexColor("#222222")
COR_SUBTEXTO   = colors.HexColor("#555555")

# Equivalentes RGB para DOCX
RGB_PRIMARIA   = RGBColor(0x1a, 0x3c, 0x5e)
RGB_SECUNDARIA = RGBColor(0x2e, 0x75, 0xb6)
RGB_FUNDO      = RGBColor(0xf0, 0xf5, 0xfa)
RGB_TEXTO      = RGBColor(0x22, 0x22, 0x22)


class GeradorCurriculo:
    def __init__(self, dados: dict, foto_bytes: bytes | None, tamanho_foto: tuple | None):
        self.dados = dados
        self.foto_bytes = foto_bytes
        self.tamanho_foto = tamanho_foto or (113, 142)  # padrão 3x4

    # ══════════════════════════════════════════════════════════════════════════
    # PDF
    # ══════════════════════════════════════════════════════════════════════════
    def gerar_pdf(self, para_impressao: bool = False) -> bytes:
        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=2*cm, rightMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm
        )
        story = self._construir_pdf(para_impressao)
        doc.build(story)
        return buf.getvalue()

    def _estilo(self):
        """Retorna dicionário de estilos customizados para PDF."""
        base = getSampleStyleSheet()
        return {
            "nome": ParagraphStyle(
                "nome", fontName="Helvetica-Bold", fontSize=22,
                textColor=COR_PRIMARIA, spaceAfter=2, leading=26
            ),
            "subtitulo": ParagraphStyle(
                "subtitulo", fontName="Helvetica", fontSize=10,
                textColor=COR_SUBTEXTO, spaceAfter=4
            ),
            "secao": ParagraphStyle(
                "secao", fontName="Helvetica-Bold", fontSize=11,
                textColor=COR_PRIMARIA, spaceBefore=10, spaceAfter=3,
                leading=14
            ),
            "corpo": ParagraphStyle(
                "corpo", fontName="Helvetica", fontSize=10,
                textColor=COR_TEXTO, spaceAfter=4, leading=14,
                alignment=TA_JUSTIFY
            ),
            "bullet": ParagraphStyle(
                "bullet", fontName="Helvetica", fontSize=10,
                textColor=COR_TEXTO, leftIndent=12, spaceAfter=2,
                leading=14, bulletIndent=0
            ),
            "cargo": ParagraphStyle(
                "cargo", fontName="Helvetica-Bold", fontSize=10,
                textColor=COR_SECUNDARIA, spaceAfter=2, leading=13
            ),
        }

    def _linha(self):
        return HRFlowable(width="100%", thickness=1, color=COR_SECUNDARIA, spaceAfter=4, spaceBefore=2)

    def _construir_pdf(self, para_impressao):
        E = self._estilo()
        D = self.dados
        story = []

        # ── Cabeçalho ──────────────────────────────────────────────────────
        foto_flowable = self._foto_pdf()

        cabecalho_texto = [
            Paragraph(D.get("nome", ""), E["nome"]),
        ]
        contato_itens = []
        if D.get("telefone"):  contato_itens.append(f"📞 {D['telefone']}")
        if D.get("email"):     contato_itens.append(f"✉ {D['email']}")
        if D.get("endereco"):  contato_itens.append(f"🏠 {D['endereco']}")
        if D.get("nascimento"):contato_itens.append(f"🎂 {D['nascimento']}")
        if D.get("cnh") and D["cnh"].lower() not in ("não", "nao", "n"):
            contato_itens.append(f"🚗 CNH: {D['cnh']}")

        for item in contato_itens:
            cabecalho_texto.append(Paragraph(item, E["subtitulo"]))

        if foto_flowable:
            tabela_header = Table(
                [[foto_flowable, cabecalho_texto]],
                colWidths=[self.tamanho_foto[0] * 0.0264 * cm * 28.35 / 28.35, None]
            )
            w_foto_pt = self.tamanho_foto[0] * (72 / 37.8)  # px → pt approx
            tabela_header = Table(
                [[foto_flowable, cabecalho_texto]],
                colWidths=[w_foto_pt + 10, None]
            )
            tabela_header.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (0, 0), 0),
                ("RIGHTPADDING", (0, 0), (0, 0), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
            ]))
            story.append(tabela_header)
        else:
            for elem in cabecalho_texto:
                story.append(elem)

        story.append(Spacer(1, 6))
        story.append(self._linha())

        # ── Resumo Profissional ─────────────────────────────────────────────
        if D.get("resumo"):
            story.append(Paragraph("RESUMO PROFISSIONAL", E["secao"]))
            story.append(self._linha())
            story.append(Paragraph(D["resumo"], E["corpo"]))

        # ── Formação Acadêmica ──────────────────────────────────────────────
        if D.get("formacao"):
            story.append(Spacer(1, 4))
            story.append(Paragraph("FORMAÇÃO ACADÊMICA", E["secao"]))
            story.append(self._linha())
            for linha in D["formacao"].split("\n"):
                if linha.strip():
                    story.append(Paragraph(linha.strip(), E["corpo"]))

        # ── Experiência Profissional ────────────────────────────────────────
        exp_texto = D.get("experiencia", "")
        if exp_texto and "sem experiência" not in exp_texto.lower():
            story.append(Spacer(1, 4))
            story.append(Paragraph("EXPERIÊNCIA PROFISSIONAL", E["secao"]))
            story.append(self._linha())
            for linha in exp_texto.split("\n"):
                linha = linha.strip()
                if not linha:
                    continue
                if linha.startswith("•") or linha.startswith("-"):
                    story.append(Paragraph(f"• {linha.lstrip('•- ').strip()}", E["bullet"]))
                else:
                    story.append(Paragraph(linha, E["cargo"]))

        # ── Competências ────────────────────────────────────────────────────
        if D.get("competencias"):
            story.append(Spacer(1, 4))
            story.append(Paragraph("COMPETÊNCIAS", E["secao"]))
            story.append(self._linha())
            for linha in D["competencias"].split("\n"):
                if linha.strip():
                    story.append(Paragraph(f"• {linha.strip()}", E["bullet"]))

        return story

    def _foto_pdf(self):
        if not self.foto_bytes:
            return None
        try:
            img = Image.open(io.BytesIO(self.foto_bytes)).convert("RGB")
            w_px, h_px = self.tamanho_foto
            img = img.resize((w_px, h_px), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="JPEG")
            buf.seek(0)
            # Converter px → pontos (72 dpi / 96 dpi padrão tela)
            w_pt = w_px * 72 / 96
            h_pt = h_px * 72 / 96
            return RLImage(ImageReader(buf), width=w_pt, height=h_pt)
        except Exception:
            return None

    # ══════════════════════════════════════════════════════════════════════════
    # DOCX
    # ══════════════════════════════════════════════════════════════════════════
    def gerar_docx(self) -> bytes:
        doc = Document()
        D = self.dados

        # Margens
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2)
            section.right_margin  = Cm(2)

        # ── Cabeçalho ──────────────────────────────────────────────────────
        if self.foto_bytes:
            self._adicionar_header_com_foto_docx(doc, D)
        else:
            self._adicionar_header_sem_foto_docx(doc, D)

        self._linha_docx(doc, RGB_SECUNDARIA)

        # ── Seções ─────────────────────────────────────────────────────────
        if D.get("resumo"):
            self._secao_titulo_docx(doc, "RESUMO PROFISSIONAL")
            self._linha_docx(doc)
            p = doc.add_paragraph(D["resumo"])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._set_fonte(p, 10)

        if D.get("formacao"):
            self._secao_titulo_docx(doc, "FORMAÇÃO ACADÊMICA")
            self._linha_docx(doc)
            for linha in D["formacao"].split("\n"):
                if linha.strip():
                    p = doc.add_paragraph(linha.strip())
                    self._set_fonte(p, 10)

        exp_texto = D.get("experiencia", "")
        if exp_texto and "sem experiência" not in exp_texto.lower():
            self._secao_titulo_docx(doc, "EXPERIÊNCIA PROFISSIONAL")
            self._linha_docx(doc)
            for linha in exp_texto.split("\n"):
                linha = linha.strip()
                if not linha:
                    continue
                if linha.startswith("•") or linha.startswith("-"):
                    p = doc.add_paragraph(style="List Bullet")
                    p.text = linha.lstrip("•- ").strip()
                    self._set_fonte(p, 10)
                else:
                    p = doc.add_paragraph(linha)
                    run = p.runs[0] if p.runs else p.add_run(linha)
                    run.bold = True
                    run.font.color.rgb = RGB_SECUNDARIA
                    self._set_fonte(p, 10)

        if D.get("competencias"):
            self._secao_titulo_docx(doc, "COMPETÊNCIAS")
            self._linha_docx(doc)
            for linha in D["competencias"].split("\n"):
                if linha.strip():
                    p = doc.add_paragraph(style="List Bullet")
                    p.text = linha.strip()
                    self._set_fonte(p, 10)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _adicionar_header_sem_foto_docx(self, doc, D):
        p = doc.add_paragraph()
        run = p.add_run(D.get("nome", ""))
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGB_PRIMARIA

        contato = []
        if D.get("telefone"):  contato.append(f"📞 {D['telefone']}")
        if D.get("email"):     contato.append(f"✉ {D['email']}")
        if D.get("endereco"):  contato.append(f"🏠 {D['endereco']}")
        if D.get("nascimento"):contato.append(f"🎂 {D['nascimento']}")
        cnh = D.get("cnh", "")
        if cnh and cnh.lower() not in ("não", "nao", "n"):
            contato.append(f"🚗 CNH: {cnh}")

        for item in contato:
            p2 = doc.add_paragraph(item)
            self._set_fonte(p2, 10, cor=RGB_TEXTO)

    def _adicionar_header_com_foto_docx(self, doc, D):
        """Usa tabela 2 colunas: foto | texto."""
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # Largura das colunas
        w_foto_cm = self.tamanho_foto[0] * 2.54 / 96 + 0.3
        table.columns[0].width = Cm(w_foto_cm)

        cell_foto = table.cell(0, 0)
        cell_texto = table.cell(0, 1)

        # Remove bordas
        for cell in [cell_foto, cell_texto]:
            for side in ["top", "bottom", "left", "right"]:
                self._set_cell_border(cell, **{side: {"sz": 0, "val": "none", "color": "FFFFFF"}})

        # Inserir foto
        try:
            img = Image.open(io.BytesIO(self.foto_bytes)).convert("RGB")
            w_px, h_px = self.tamanho_foto
            img = img.resize((w_px, h_px), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="JPEG")
            buf.seek(0)
            p_foto = cell_foto.paragraphs[0]
            run = p_foto.add_run()
            run.add_picture(buf, width=Cm(w_foto_cm - 0.3))
        except Exception:
            cell_foto.text = ""

        # Texto no lado direito
        p = cell_texto.paragraphs[0]
        run = p.add_run(D.get("nome", ""))
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGB_PRIMARIA

        contato = []
        if D.get("telefone"):  contato.append(f"📞 {D['telefone']}")
        if D.get("email"):     contato.append(f"✉ {D['email']}")
        if D.get("endereco"):  contato.append(f"🏠 {D['endereco']}")
        if D.get("nascimento"):contato.append(f"🎂 {D['nascimento']}")
        cnh = D.get("cnh", "")
        if cnh and cnh.lower() not in ("não", "nao", "n"):
            contato.append(f"🚗 CNH: {cnh}")

        for item in contato:
            p2 = cell_texto.add_paragraph(item)
            self._set_fonte(p2, 10)

    def _secao_titulo_docx(self, doc, titulo: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGB_PRIMARIA

    def _linha_docx(self, doc, cor: RGBColor = None):
        cor = cor or RGBColor(0xd0, 0xdc, 0xe8)
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), f"{cor[0]:02X}{cor[1]:02X}{cor[2]:02X}")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

    def _set_fonte(self, p, size: int, negrito=False, cor: RGBColor = None):
        for run in p.runs:
            run.font.size = Pt(size)
            if negrito:
                run.bold = True
            if cor:
                run.font.color.rgb = cor

    def _set_cell_border(self, cell, **kwargs):
        """Define bordas de uma célula."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge, attrs in kwargs.items():
            edge_el = OxmlElement(f"w:{edge}")
            for k, v in attrs.items():
                edge_el.set(qn(f"w:{k}"), str(v))
            tcBorders.append(edge_el)
        tcPr.append(tcBorders)
